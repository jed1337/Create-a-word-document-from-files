from path import Path
from PathHolder import PathHolder
from docx import Document


def print_all_files(path):
    if(path.isdir()):
        print(f"{path} is a directory")
        for sub_path in path.listdir():
            print_all_files(sub_path)
    elif(path.isfile()):
        print(f"{path} is a file")


def parent_and_basename(path):
    return path.parent.name+"\\"+path.name


def save_to_word(path_holders):
    document = Document()
    document.add_heading('Document Title', level=0)

    document.add_heading('Scripts', level=1)

    for path_holder in path_holders:
        document.add_paragraph(path_holder.path.name, style='List Number')
        for script_name in path_holder.script_names:
            document.add_paragraph(parent_and_basename(script_name), style='List Bullet 2')

    document.add_heading('VERIFICATION', level=1)
    for path_holder in path_holders:
        document.add_paragraph(path_holder.path.name, style='List Number')
        for verification_script_name in path_holder.verification_script_names:
            document.add_paragraph(parent_and_basename(verification_script_name), style='List Bullet 2')

    document.save('test.docx')

path = Path(r"D:\path")
path_holders = [PathHolder(sub_path) for sub_path in path.listdir() if sub_path.isdir()]
save_to_word(path_holders)
